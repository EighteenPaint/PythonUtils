# -*- coding:utf-8 -*-
'''
Created by Administrator on 2019/3/14 0014.
'''
#目前发现的类型：NUMBER，DATE，VARCHAR2
import re
# 提取字典数据
def extractor_dict_data(file_name):
    from docx import Document
    doc = Document(file_name)  # type:docx.table._Row
    for index,tables in enumerate(doc.tables):
        type_code = index + 1
        for row in tables.rows:
            cells = row.cells
            if cells[0].text.strip() != '':
                code = cells[0].text.strip()
                name = cells[1].text.strip()
                if cells[2].text.strip() != u'说明':
                 print u"('{}', '{}' ,'{}' ,{}) ,".format(code,name,cells[2].text,type_code)
def extractor_jq_data():
    from docx import Document
    doc = Document("FKDB.docx")  # type:docx.table._Row
    count = 0
    for tables in doc.tables:
        for row in tables.rows:
            cells = row.cells
            null_able = ''
            data_type = 'error'
            if cells[2].text == u'N':
                null_able = u'NOT NULL'
            matchObj = re.match(r'(.*)\((.*)\)|(DATE)', cells[1].text, re.M | re.I)
            if matchObj:
                if matchObj.group(0).strip() == 'DATE':
                    # print u'DATE'
                    data_type = u'DATETIME'
                else:
                    # print u'{} {}'.format(matchObj.group(1),matchObj.group(2))
                    if matchObj.group(1).strip() == u'VARCHAR2' or matchObj.group(1).strip() == u'VARCHAR':
                        if int(matchObj.group(2)) > 255:
                            data_type = u'TEXT'
                        else:
                            data_type = u'VARCHAR({})'.format(matchObj.group(2))
                    if matchObj.group(1).strip() == u'NUMBER':
                        if re.match(r'.*,.*',matchObj.group(0)):
                            data_type = u'DECIMAL({})'.format(matchObj.group(2).strip())
                        else:
                            data_type = u'INT(11)'  # .format(matchObj.group(2))
                    if matchObj.group(1).strip() == u'DATE':
                        data_type = u'DATETIME'

            count = count + 1
            print u"`{}` {} {} COMMENT '{}',\n".format(cells[0].text,data_type,null_able,cells[3].text)
            # 警情数据协议体
            # print u'\"{}\":____,//{}'.format(cells[0].text, cells[3].text)
if __name__ == '__main__':
 # extractor_dict_data(u"DICT_JJLB.docx")
 extractor_jq_data()
